from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(__file__).parent / "Verixora_Technical_Assessment_and_Product_Blueprint.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    paragraph.add_run("Page ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("VERIXORA  |  TECHNICAL ASSESSMENT & PRODUCT BLUEPRINT")
    set_font(r, size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Confidential - Internal planning document  |  ")
    set_font(r, size=8, color=MUTED)
    add_page_field(footer)


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("VERIXORA")
    set_font(r, size=25, color=INK, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(15)
    r = p.add_run("Technical Assessment & Product Blueprint")
    set_font(r, size=15, color=DARK_BLUE, bold=True)
    for label, value in [
        ("Purpose", "Current codebase documentation and production-readiness blueprint"),
        ("Scope", "API host, .NET modules, security model, ESP32/MQTT integration, and face service"),
        ("Status", "Architecture assessment - no implementation changes proposed in this document"),
        ("Prepared", "19 July 2026"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        a = p.add_run(f"{label}: ")
        set_font(a, size=10.5, color=INK, bold=True)
        b = p.add_run(value)
        set_font(b, size=10.5, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Executive summary")
    set_font(r, size=13, color=BLUE, bold=True)
    add_para(doc, "Verixora is an early-stage modular .NET 8 smart-access backend with a separate Python face-verification service. The repository already proves the main technical direction: identity, devices, smart locks, audit logs, MQTT publishing, and face verification. It builds successfully, but several controls are currently demo-level and must be strengthened before real users or physical locks are connected.")
    add_callout(doc, "Key conclusion", "The best next step is to treat this codebase as a foundation: first harden identity and trusted-device binding, then secure ESP32 provisioning and MQTT, then add proximity, face liveness, operational monitoring, and production deployment controls.")
    doc.add_page_break()


def add_para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_font(r, bold=True)
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title.upper())
    set_font(r, size=9, color=DARK_BLUE, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, value in zip(header.cells, headers):
        shade(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_font(r, size=9.5, color=DARK_BLUE, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_font(r, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def build():
    doc = Document()
    style_document(doc)
    add_title(doc)

    doc.add_heading("1. Repository at a glance", level=1)
    add_para(doc, "The solution is a modular monolith. ApiHost composes independently structured modules using dependency injection and MediatR. Each active module generally follows Domain, Application, Infrastructure, and Presentation layers.")
    add_table(doc, ["Area", "Current responsibility", "Status"], [
        ("ApiHost", "ASP.NET Core host, JWT middleware, Swagger, module composition", "Active"),
        ("Identity", "Phone/password/OTP registration and login; email verification; JWT", "Active"),
        ("Devices", "IoT device registration and MQTT topic generation", "Active"),
        ("SmartLocks", "Lock registration, authorization, face verification, unlock command", "Active"),
        ("AuditLogs", "Records actions and exposes logs by home", "Active"),
        ("Homes", "Home and membership domain model", "Domain-only"),
        ("Face service", "FastAPI enrollment and verification endpoint", "Active, demo persistence"),
    ], [1900, 5200, 2260])

    doc.add_heading("2. Current architecture", level=1)
    add_para(doc, "ApiHost is the composition root. It configures SQLite by default, JWT bearer authentication, authorization middleware, Swagger in development, MediatR handler discovery, and ASP.NET controllers from the module presentation assemblies.")
    add_para(doc, "The application uses domain entities and repository interfaces in the Application layer, with Entity Framework Core repositories in Infrastructure. This separation is a sound base for a future production system.")
    add_table(doc, ["Layer", "Purpose", "Examples"], [
        ("Domain", "Business entities, value rules, status enums, domain events", "User, SmartLock, Device, Home, AuditLog"),
        ("Application", "Commands, queries, interfaces, use-case handlers", "UnlockDoorCommandHandler, LoginCommandHandler"),
        ("Infrastructure", "EF Core contexts/repositories and external adapters", "JwtTokenGenerator, MqttPublisher, PythonFaceVerificationProvider"),
        ("Presentation", "HTTP controllers and request binding", "AuthController, DevicesController, SmartLocksController"),
    ], [1500, 3900, 3960])

    doc.add_heading("3. Implemented API surface", level=1)
    add_table(doc, ["Route", "Purpose", "Notes"], [
        ("POST /api/v1/auth/send-otp", "Send registration OTP", "Currently backed by a mock OTP service"),
        ("POST /api/v1/auth/register", "Register phone/password/OTP user", "Optional email can be supplied"),
        ("POST /api/v1/auth/send-login-otp", "Start phone login OTP", "Password is included in command"),
        ("POST /api/v1/auth/login", "Phone/password/OTP login", "Returns a JWT"),
        ("POST /api/v1/auth/set-email", "Set user email", "Uses user ID supplied in request"),
        ("POST /api/v1/auth/send-verification-email", "Send email verification code", "Mock email/OTP service"),
        ("POST /api/v1/auth/verify-email", "Verify email", "Enables web login eligibility"),
        ("POST /api/v1/auth/web/send-login-otp", "Start email-based web OTP", "Non-guest, verified email path"),
        ("POST /api/v1/auth/web/login", "Email/password/OTP web login", "Returns a JWT"),
        ("POST /api/v1/devices", "Register an IoT device", "Immediately activates device for demo"),
        ("POST /api/v1/locks", "Register a smart lock", "Links device ID and home ID"),
        ("POST /api/v1/locks/{id}/unlock", "Unlock with multipart face image", "Runs authorization, optional face check, MQTT, audit"),
        ("GET /api/v1/auditlogs?homeId=", "Read audit logs for a home", "No access boundary currently enforced"),
    ], [2850, 3100, 3410])

    doc.add_heading("4. Current unlock workflow", level=1)
    add_numbers(doc, [
        "The client calls the unlock route with a lock ID, face image when required, and an optional idempotency key.",
        "The handler loads the SmartLock aggregate from the repository.",
        "ScheduleBasedAuthorizationService evaluates the role: Owner is allowed; Guest is allowed only on weekdays from 09:00 to 17:00 UTC.",
        "If RequiresFace is true, PythonFaceVerificationProvider sends the image and user ID to the FastAPI /verify endpoint.",
        "The lock aggregate changes to Unlocked and stores LastUnlockedAt and LastUnlockedBy.",
        "The host publishes an MQTT unlock payload to verixora/{deviceId}.",
        "AuditLogs records a successful or rejected attempt.",
    ])
    add_callout(doc, "Observed behavior", "The domain model correctly captures a lock state change and domain event. However, the current flow does not wait for a physical ESP32 acknowledgement that the relay actuated, nor does it establish that the requester is near the door.")

    doc.add_heading("5. Identity and roles currently in code", level=1)
    add_para(doc, "A User has a phone number, password hash, optional email, email-verification flag, role, one TrustedDevice navigation property, and a collection of FaceEmbedding records. The initial user role is Owner. The existing UserRole enum defines Owner and Guest.")
    add_para(doc, "The code includes a TrustedDevice entity and domain events, but the registration/login handlers do not currently create, bind, validate, or revoke trusted devices. Device binding therefore remains a planned capability rather than an enforced security boundary.")

    doc.add_heading("6. Data persistence and integrations", level=1)
    add_bullets(doc, [
        "SQLite is the configured default database. The modules use separate EF Core DbContexts and schemas: identity, devices, smartlocks, and auditlogs.",
        "SmartLocks can use either EF Core or an ADO.NET repository depending on UseEfCore configuration.",
        "MQTT uses a singleton publisher and defaults to localhost:1883.",
        "The Python face service uses FastAPI, Pillow, NumPy, and face_recognition. Embeddings are currently retained only in process memory.",
        "The built solution compiles successfully with zero warnings and zero errors when built without restore.",
    ])

    doc.add_page_break()
    doc.add_heading("7. Production readiness gaps", level=1)
    add_para(doc, "The following are implementation observations, not criticism of the architecture. They identify the work required before connecting physical locks or onboarding real users.")
    add_table(doc, ["Priority", "Gap", "Why it matters", "Recommended direction"], [
        ("Critical", "No per-endpoint authorization attributes/policies", "An unlock controller can obtain empty/Guest claims for an unauthenticated caller.", "Require authenticated policies by default; enforce ownership/permission in handlers."),
        ("Critical", "Trusted-device model is not enforced", "A token alone does not meet the single-mobile-device requirement.", "Use hardware-backed keys, attestation, device signatures, and server-side binding."),
        ("Critical", "Mock OTP and email implementations", "Cannot verify real phone/email ownership.", "Integrate a production SMS/email provider with expiry, rate limits, and fraud controls."),
        ("Critical", "Face embeddings held in memory", "Data is lost on restart and is not privacy/security ready.", "Encrypted persistence, consent handling, liveness detection, and retention policies."),
        ("High", "Static development JWT secret in configuration", "Secrets must not be committed or shared across environments.", "Use environment-specific secret management and key rotation."),
        ("High", "No ESP32 acknowledgement / command replay defense", "MQTT publish does not prove the door opened safely.", "Per-device certificates, signed short-lived commands, nonce, command IDs, and acknowledgements."),
        ("High", "Idempotency key is not consumed", "Duplicate unlock attempts can result in duplicate commands/logs.", "Persist and enforce idempotency per user/device/lock."),
        ("High", "Homes module not wired", "Ownership, membership, and tenant boundaries cannot be managed yet.", "Complete Home application, infrastructure, API, and authorization policies."),
        ("Medium", "No startup migration/init process", "Fresh installations may lack the required schema.", "Controlled migrations in CI/CD and a deployment-safe database lifecycle."),
        ("Medium", "No evidence of tests/observability", "Physical access systems need operational confidence and incident traces.", "Unit/integration tests, metrics, alerts, structured logs, and backups."),
    ], [900, 1800, 3000, 3660])

    doc.add_heading("8. Recommended product model", level=1)
    add_para(doc, "To support homes today and buildings, offices, or multiple client deployments later, use an explicit tenant hierarchy rather than directly attaching everything to a single user.")
    add_callout(doc, "Recommended hierarchy", "SuperAdmin -> Client / Organization -> Site / Home -> Door -> ESP32 Controller -> Access policies, users, credentials, and events.")
    add_table(doc, ["Role", "Suggested capability boundary"], [
        ("SuperAdmin", "Global tenant, device, audit, support, firmware, and system administration."),
        ("ClientAdmin", "Own organization/sites, doors, members, schedules, and reports."),
        ("Manager / Operator", "Assigned doors, member operations, and local audit review."),
        ("Resident", "Use only specifically granted doors and settings."),
        ("Guest", "Temporary, schedule-limited, door-limited access only."),
        ("Installer", "Time-limited device provisioning and installation diagnostics."),
    ], [2100, 7260])

    doc.add_heading("9. Secure mobile registration and device binding", level=1)
    add_para(doc, "The requirement should be interpreted as one active trusted application/device binding per user, not as a permanent and infallible browser-style device fingerprint. Operating systems deliberately delete app storage on uninstall, and a device fingerprint alone can be copied or changed.")
    add_numbers(doc, [
        "Mobile app validates the phone number using a country-aware phone-number library and submits it with device attestation and a newly generated public key.",
        "Backend rate-limits the request and sends an SMS OTP through a real provider.",
        "User submits the OTP. Backend creates the account and one active TrustedDevice record only if the device is not already bound to another active identity.",
        "The mobile app creates the private key inside Android Keystore / iOS Secure Enclave or equivalent hardware-backed storage. The private key never leaves the device.",
        "Every sensitive request includes a short-lived access token plus a device-signed proof of request. Backend validates the signature, device status, key version, and attestation risk signal.",
        "On uninstall/reinstall, create a new binding after strong re-verification. Do not silently trust a recreated identifier; this is a security control.",
    ])
    add_callout(doc, "Important policy", "For real security, new-device activation, device replacement, root/jailbreak detection, and recovery must be explicit business flows with notification and audit logging.")

    doc.add_heading("10. Email verification and web portal", level=1)
    add_para(doc, "The planned web portal should remain a separate access channel with stronger verification. A user sets email in the trusted mobile app, verifies an email OTP/link, and only then gains eligibility for web login. Web login should require email/password plus email OTP, and can add passkeys or authenticator-app MFA later.")
    add_bullets(doc, [
        "Only roles with web entitlement can enter the portal; Guest has no portal access.",
        "The API must derive user identity from the authenticated token, never from a user ID posted by the client.",
        "Settings should include device management, email management, password change, face enrollment/re-enrollment, notification preferences, active sessions, recovery, and activity history.",
    ])

    doc.add_page_break()
    doc.add_heading("11. ESP32 provisioning and secure MQTT", level=1)
    add_para(doc, "Each ESP32 DevKit V1 should be assigned to exactly one Door at a time. The backend must prevent an unclaimed or already-claimed controller from being registered by an arbitrary app user.")
    add_table(doc, ["Provisioning control", "Recommended design"], [
        ("Device identity", "Unique factory device ID and device-specific secret or certificate. Never rely only on a printed MAC address."),
        ("Claiming", "Installer/SuperAdmin scans a QR code and claims the controller during short pairing mode."),
        ("Door binding", "Backend binds one controller to one door. Moving it requires a deliberate unassign/reassign workflow and audit trail."),
        ("Transport", "MQTT over TLS, per-device credentials/certificates, least-privilege topic ACLs."),
        ("Commands", "Signed, short-lived unlock command containing command ID, nonce, expiry, target door, and authorization context."),
        ("Acknowledgement", "ESP32 reports accepted/rejected, relay action, door sensor state, and error code. Backend finalizes the audit event after acknowledgement."),
        ("Firmware", "Signed OTA updates, version inventory, rollback policy, and firmware health telemetry."),
    ], [2450, 6910])
    add_para(doc, "Recommended sensors and operational signals include a reed switch for door open/closed state, lock/relay feedback when hardware supports it, tamper switch, power/battery status, Wi-Fi health, firmware version, and optional PIR motion telemetry. Sensor data should be treated as operational evidence, not as the sole authorization factor.")

    doc.add_heading("12. Presence and face verification", level=1)
    add_para(doc, "GPS alone is unsuitable for secure indoor door presence. A better approach is a near-door Bluetooth Low Energy (BLE) challenge-response protocol. The ESP32 broadcasts a short-lived challenge; only an authenticated nearby app can complete a server-authorized, signed response before the unlock request proceeds.")
    add_numbers(doc, [
        "Confirm a trusted mobile device and valid session.",
        "Confirm the user's access policy for that exact door and time window.",
        "Complete BLE near-door challenge-response.",
        "Confirm face enrollment exists; otherwise reject unlock before calling the controller.",
        "Perform liveness detection and face match using encrypted biometric templates.",
        "Issue the signed ESP32 command and wait for acknowledgement.",
        "Write tamper-evident audit events for each decision and final physical result.",
    ])
    add_callout(doc, "Biometric privacy", "Face templates are sensitive personal data. Collect explicit consent, encrypt templates at rest, restrict access, define retention/deletion rules, and never treat a simple image comparison as sufficient liveness protection.")

    doc.add_heading("13. Implementation roadmap", level=1)
    add_table(doc, ["Phase", "Outcome", "Main deliverables"], [
        ("1. Foundation", "Secure API foundation", "Global auth policy, tenant/home model, database migrations, secrets, error handling, test baseline."),
        ("2. Identity", "Phone and trusted-device security", "Real SMS OTP, phone validation, hardware key binding, attestation, refresh/revocation, recovery."),
        ("3. Hardware", "Secure controller lifecycle", "Device registry, QR claim flow, TLS MQTT, command/ack protocol, ESP32 firmware baseline."),
        ("4. Access", "Reliable unlock authorization", "Door permissions, schedules, BLE presence, idempotency, physical-state audit events."),
        ("5. Biometrics", "Privacy-safe face access", "Enrollment, liveness, encrypted template storage, re-enrollment and deletion controls."),
        ("6. Operations", "Commercial readiness", "Web portal, dashboards, tenant isolation, monitoring, alerts, backups, incident tooling."),
        ("7. Launch", "Controlled production deployment", "Threat model review, penetration testing, pilot hardware rollout, runbooks, support process."),
    ], [1400, 2500, 5460])

    doc.add_heading("14. Recommended next design decisions", level=1)
    add_numbers(doc, [
        "Confirm the business tenant model: single-home customers only, or organizations with many sites and doors from day one.",
        "Choose the mobile platform approach: native Android/iOS or a cross-platform app that still supports secure hardware keystore and platform attestation.",
        "Choose an SMS provider, email provider, cloud/deployment target, and managed MQTT broker approach for the first production environment.",
        "Select the physical lock, relay/driver, power supply, enclosure, door sensor, and optional BLE/sensor hardware before finalizing firmware interfaces.",
        "Create a written threat model covering stolen phones, cloned/replayed requests, compromised ESP32 devices, lost connectivity, biometric spoofing, and administrator misuse.",
    ])

    doc.add_heading("Appendix A. Source files reviewed", level=1)
    add_para(doc, "This assessment was derived from the repository structure and core implementation files, including ApiHost Program.cs and configuration; Identity, Devices, SmartLocks, AuditLogs, and Homes projects; EF Core contexts and dependency-injection registrations; and face-service/app.py. This document deliberately distinguishes existing behavior from recommended future architecture.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
